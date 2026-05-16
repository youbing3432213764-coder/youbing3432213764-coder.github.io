"""
Batch convert docx files from 题解/ to Hexo markdown posts.
Extracts text, images, and generates proper frontmatter.
"""
import os
import re
import zipfile
import shutil
import hashlib
from pathlib import Path
from datetime import datetime

BASE_DIR = Path(r"D:\blog\youbingblog")
SOURCE_DIR = BASE_DIR / "题解"
POSTS_DIR = BASE_DIR / "source" / "_posts"
IMAGES_DIR = BASE_DIR / "source" / "images" / "writeups"
SKIP_PREFIX = "甄洪琳的"
SKIP_SUFFIX = "题解"
SKIP_SUFFIX2 = "题解·"

def sanitize_filename(name):
    """Remove characters that are problematic in filenames."""
    name = name.replace("？", "").replace("！", "").replace("!", "").replace("？", "")
    name = name.replace("：", "").replace("（", "(").replace("）", ")")
    name = name.replace("。", "").replace("，", "")
    name = re.sub(r'[\\/*?"<>|]', '', name)
    name = name.strip()
    name = re.sub(r'\s+', '-', name)
    return name

def extract_title(filename):
    """Extract a clean title from the docx filename."""
    stem = Path(filename).stem
    # Remove "甄洪琳的" prefix
    if stem.startswith(SKIP_PREFIX):
        stem = stem[len(SKIP_PREFIX):]
    # Remove "甄洪琳" suffix (without "的")
    if stem.endswith("甄洪琳"):
        stem = stem[:-3]
    # Remove "题解" suffix
    stem = re.sub(r'题解[·.]?$', '', stem)
    # Remove leading/trailing whitespace
    stem = stem.strip()
    return stem

def get_date_from_dir(dirname):
    """Extract date from directory name like '1.25', '2.3', etc."""
    if re.match(r'^\d+\.\d+$', dirname):
        parts = dirname.split('.')
        month = int(parts[0])
        day = int(parts[1])
        return f"2026-{month:02d}-{day:02d}"
    return None

def extract_images_from_docx(docx_path, output_dir, prefix):
    """Extract embedded images from a docx file."""
    image_map = {}  # rId -> output_filename
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            # Find all image files in the docx
            image_files = [f for f in z.namelist() if f.startswith('word/media/')]

            for img_file in image_files:
                img_data = z.read(img_file)
                # Generate unique filename based on content hash
                content_hash = hashlib.md5(img_data).hexdigest()[:8]
                ext = Path(img_file).suffix
                if not ext:
                    ext = '.png'
                new_name = f"{prefix}_{content_hash}{ext}"
                new_path = output_dir / new_name

                # Only write if not already exists
                if not new_path.exists():
                    with open(new_path, 'wb') as f:
                        f.write(img_data)

                # Map the relationship filename
                img_relname = img_file.replace('word/', '')
                image_map[img_relname] = f"/images/writeups/{new_name}"
    except Exception as e:
        print(f"  Warning extracting images: {e}")

    return image_map

def get_tag_from_title(title):
    """Determine tags based on title content."""
    title_lower = title.lower()
    tags = ["题解"]

    # Common CTF/web categories
    if any(kw in title_lower for kw in ['sql', '注入', 'sqli', '盲注']):
        tags.append('SQL注入')
    if any(kw in title_lower for kw in ['xss', '跨站']):
        tags.append('XSS')
    if any(kw in title_lower for kw in ['upload', '上传', '文件上传']):
        tags.append('文件上传')
    if any(kw in title_lower for kw in ['rce', '命令执行', '代码执行', 'eval', 'ping', 'exec']):
        tags.append('RCE')
    if any(kw in title_lower for kw in ['ssti', '模板注入']):
        tags.append('SSTI')
    if any(kw in title_lower for kw in ['反序列化', 'serialize', 'unser', 'unse', 'pop']):
        tags.append('反序列化')
    if any(kw in title_lower for kw in ['include', '包含', '文件包含', 'lfi', 'rfi']):
        tags.append('文件包含')
    if any(kw in title_lower for kw in ['php', '躲猫猫', '特性']):
        tags.append('PHP')
    if any(kw in title_lower for kw in ['jwt', 'jwt']):
        tags.append('JWT')
    if any(kw in title_lower for kw in ['session', 'cookie', '伪造']):
        tags.append('Session')
    if any(kw in title_lower for kw in ['base64', '编码', '解码']):
        tags.append('编码解码')
    if any(kw in title_lower for kw in ['爆破', 'password', 'weak', '弱口令']):
        tags.append('爆破')
    if any(kw in title_lower for kw in ['ff', 'xff', 'ip', '头']):
        tags.append('HTTP头')
    if any(kw in title_lower for kw in ['shell', '马']):
        tags.append('Webshell')
    if any(kw in title_lower for kw in ['gif', '二维码', '照片', '美照', '图片', '羽毛球']):
        tags.append('图片隐写')
    if any(kw in title_lower for kw in ['zip', '压缩', '压缩包']):
        tags.append('压缩包')
    if any(kw in title_lower for kw in ['flask', 'python', 'flask']):
        tags.append('Flask')
    if any(kw in title_lower for kw in ['ping', 'pingping']):
        tags.append('命令注入')
    if any(kw in title_lower for kw in ['源码', '审计', '代码']):
        tags.append('代码审计')
    if any(kw in title_lower for kw in ['cms', 'easycms']):
        tags.append('CMS')

    return tags

def convert_paragraph_to_markdown(para, image_map, image_index):
    """Convert a paragraph element to markdown string.
    Returns (markdown_text, new_image_index)"""
    style = para.style.name if para.style else 'Normal'
    text_parts = []

    for run in para.runs:
        # Check for images in this run
        drawings = run._element.findall(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
        )
        for drawing in drawings:
            blips = drawing.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
            )
            for blip in blips:
                embed = blip.get(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                )
                if embed:
                    # Look up in the document's rels
                    try:
                        rel = para.part.rels[embed]
                        rel_ref = rel.target_ref
                        if rel_ref in image_map:
                            text_parts.append(f"\n![image{image_index}]({image_map[rel_ref]})\n")
                            image_index += 1
                        else:
                            # Try partial match
                            for key, val in image_map.items():
                                if rel_ref.endswith(key) or key.endswith(rel_ref):
                                    text_parts.append(f"\n![image{image_index}]({val})\n")
                                    image_index += 1
                                    break
                    except (KeyError, AttributeError):
                        pass

        # Also check for inline shapes (alternative image embedding)
        drawings2 = run._element.findall(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'
        )
        # Skip the pict elements - they're handled above

        text = run.text
        if text:
            bold = run.bold
            if bold:
                text_parts.append(f"**{text}**")
            else:
                text_parts.append(text)

    result = ''.join(text_parts)

    # Handle heading styles
    if 'Heading 1' in style or style == 'heading 1':
        result = f"# {result}" if result.strip() else ""
    elif 'Heading 2' in style or style == 'heading 2':
        result = f"## {result}" if result.strip() else ""
    elif 'Heading 3' in style or style == 'heading 3':
        result = f"### {result}" if result.strip() else ""

    return result, image_index

def convert_paragraph_to_markdown_simple(para, image_map, used_images):
    """Simpler conversion that handles most cases."""
    text = para.text
    if not text and not para.runs:
        return ""

    style = para.style.name if para.style else 'Normal'

    # Check for inline images
    result_parts = []
    for run in para.runs:
        # Check drawings
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        }

        # Find all blip elements (image references)
        blips = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        for blip in blips:
            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if embed and embed in image_map:
                img_path = image_map[embed]
                if img_path not in used_images:
                    used_images.add(img_path)
                    result_parts.append(f"\n![{Path(img_path).stem}]({img_path})\n")

        if run.text:
            if run.bold:
                result_parts.append(f"**{run.text}**")
            else:
                result_parts.append(run.text)

    result = ''.join(result_parts)

    # Handle headings
    if style in ('Heading 1', 'heading 1'):
        result = f"# {result}" if result.strip() else ""
    elif style in ('Heading 2', 'heading 2'):
        result = f"## {result}" if result.strip() else ""
    elif style in ('Heading 3', 'heading 3'):
        result = f"### {result}" if result.strip() else ""
    elif style in ('HTML Preformatted',):
        # For code or preformatted blocks
        if '\n' in text:
            result = f"```\n{text}\n```"
        else:
            result = f"`{text}`"

    return result

def convert_docx_to_md(docx_path, output_md_path, title, date_str):
    """Convert a single docx file to markdown."""
    try:
        from docx import Document

        doc = Document(str(docx_path))

        # Build image map from document relationships
        image_map = {}
        for rId, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                rel_ref = rel.target_ref
                if rel_ref.startswith('media/'):
                    img_name = Path(rel_ref).stem
                    # Search for this image in our output directory
                    for f in IMAGES_DIR.glob(f"{title[:20]}*"):
                        fname = f.name
                        # Match based on the relId or similar
                        pass
                    image_map[rId] = None  # Will be filled later

        # Actually use zipfile to map images to files
        zip_image_map = extract_images_from_docx(
            str(docx_path),
            IMAGES_DIR,
            sanitize_filename(title)[:30]
        )

        # Build proper image map by matching zip contents to rels
        for rId, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                rel_ref = rel.target_ref
                for zref, zpath in zip_image_map.items():
                    if rel_ref.endswith(zref) or zref.endswith(rel_ref) or rel_ref == zref:
                        image_map[rId] = zpath
                        break
                if rId not in image_map or image_map[rId] is None:
                    # Try matching by checking
                    for zref, zpath in zip_image_map.items():
                        image_map[rId] = zpath
                        break

        # Convert paragraphs
        markdown_lines = []
        used_images = set()

        for para in doc.paragraphs:
            md = convert_paragraph_to_markdown_simple(para, image_map, used_images)
            if md.strip():
                markdown_lines.append(md)

        # Handle tables if present
        for table in doc.tables:
            markdown_lines.append("")
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cells.append(cell.text.replace('\n', ' '))
                markdown_lines.append("| " + " | ".join(cells) + " |")
            markdown_lines.append("")

        if not markdown_lines:
            return False

        content = '\n\n'.join(markdown_lines)

        # Generate tags
        tags = get_tag_from_title(title)
        tags_str = '\n  - '.join(tags)

        # Build frontmatter
        frontmatter = f"""---
title: {title}
date: {date_str} 12:00:00
tags:
  - {tags_str}
---

"""

        full_content = frontmatter + content

        # Write markdown file
        with open(output_md_path, 'w', encoding='utf-8') as f:
            f.write(full_content)

        return True
    except Exception as e:
        print(f"  Error converting: {e}")
        return False

def main():
    os.makedirs(POSTS_DIR, exist_ok=True)
    os.makedirs(IMAGES_DIR, exist_ok=True)

    # Collect all docx files
    docx_files = list(SOURCE_DIR.rglob("*.docx"))
    # Also check for .docx files with Chinese dot (·)
    docx_files += list(SOURCE_DIR.rglob("*.docx"))

    # Deduplicate
    seen = set()
    unique_files = []
    for f in docx_files:
        if str(f) not in seen:
            seen.add(str(f))
            unique_files.append(f)

    print(f"Found {len(unique_files)} docx files")

    success = 0
    fail = 0
    skip = 0

    for docx_path in sorted(unique_files):
        rel_dir = docx_path.parent.relative_to(SOURCE_DIR)
        dirname = str(rel_dir).replace('\\', '/')
        filename = docx_path.name

        # Extract title
        title = extract_title(filename)
        if not title:
            print(f"SKIP (empty title): {filename}")
            skip += 1
            continue

        # Get date
        date_str = get_date_from_dir(docx_path.parent.name)
        if not date_str:
            # Use file modification time for non-date directories
            mtime = os.path.getmtime(str(docx_path))
            dt = datetime.fromtimestamp(mtime)
            date_str = dt.strftime("%Y-%m-%d")

        # Generate output filename
        safe_name = sanitize_filename(title)
        md_filename = f"{date_str}-{safe_name}.md"
        output_path = POSTS_DIR / md_filename

        # Skip if already converted
        if output_path.exists():
            print(f"SKIP (exists): {md_filename}")
            skip += 1
            continue

        print(f"Converting: {dirname}/{filename}")
        print(f"  -> {md_filename}")

        if convert_docx_to_md(docx_path, output_path, title, date_str):
            success += 1
        else:
            fail += 1

    print(f"\n{'='*60}")
    print(f"Done! Success: {success}, Failed: {fail}, Skipped: {skip}")

if __name__ == "__main__":
    main()
